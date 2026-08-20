import streamlit as st
import faiss
import os
from io import BytesIO
from docx import Document
import numpy as np
from PyPDF2 import PdfReader
from langchain.chains import RetrievalQA
from langchain_text_splitters import CharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.docstore.in_memory import InMemoryDocstore
from streamlit.runtime.uploaded_file_manager import UploadedFile

huggingface_api_key = st.secrets["huggingface_api_key"]

# Set the Hugging Face Hub API token as an environment variable
os.environ['HUGGINGFACEHUB_API_TOKEN'] = huggingface_api_key

def process_input(input_type, input_data):
    """Processes different input types and returns a vectorstore."""

    if input_type == "PDF":
        if input_data is None:
            raise ValueError("Please upload a PDF file")
        if isinstance(input_data, BytesIO):
            pdf_reader = PdfReader(input_data)
        elif isinstance(input_data, UploadedFile):
            pdf_reader = PdfReader(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data for PDF")
        
        text = ""
        for page in pdf_reader.pages:
            extracted = page.extract_text()
            if extracted:
                text += extracted
        
        if not text.strip():
            raise ValueError("No text could be extracted from the PDF")
        
        text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        texts = text_splitter.split_text(text)
        
    elif input_type == "Text":
        if not input_data or not input_data.strip():
            raise ValueError("Please enter some text")
        text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        texts = text_splitter.split_text(input_data)
        
    elif input_type == "DOCX":
        if input_data is None:
            raise ValueError("Please upload a DOCX file")
        if isinstance(input_data, BytesIO):
            doc = Document(input_data)
        elif isinstance(input_data, UploadedFile):
            doc = Document(BytesIO(input_data.read()))
        else:
            raise ValueError("Invalid input data for DOCX")
        
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        if not text.strip():
            raise ValueError("No text could be extracted from the DOCX file")
        
        text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        texts = text_splitter.split_text(text)
        
    elif input_type == "TXT":
        if input_data is None:
            raise ValueError("Please upload a TXT file")
        if isinstance(input_data, BytesIO):
            text = input_data.read().decode('utf-8')
        elif isinstance(input_data, UploadedFile):
            text = input_data.read().decode('utf-8')
        else:
            raise ValueError("Invalid input data for TXT")
        
        if not text.strip():
            raise ValueError("The TXT file is empty")
        
        text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
        texts = text_splitter.split_text(text)
    else:
        raise ValueError("Unsupported input type")

    # Create embeddings
    model_name = "sentence-transformers/all-mpnet-base-v2"
    model_kwargs = {'device': 'cpu'}
    encode_kwargs = {'normalize_embeddings': False}

    hf_embeddings = HuggingFaceEmbeddings(
        model_name=model_name,
        model_kwargs=model_kwargs,
        encode_kwargs=encode_kwargs
    )
    
    # Create FAISS index
    sample_embedding = np.array(hf_embeddings.embed_query("sample text"))
    dimension = sample_embedding.shape[0]
    index = faiss.IndexFlatL2(dimension)
    
    vector_store = FAISS(
        embedding_function=hf_embeddings.embed_query,
        index=index,
        docstore=InMemoryDocstore(),
        index_to_docstore_id={},
    )
    vector_store.add_texts(texts)
    return vector_store

def answer_question(vectorstore, query):
    """Answers a question using Groq API (fast and powerful!)"""
    try:
        from groq import Groq
        groq_api_key = st.secrets["groq_api_key"]
    except ImportError:
        st.error("Groq library not installed. Run: pip install groq")
        return "❌ Groq library missing"
    
    try:
        docs = vectorstore.similarity_search(query, k=4)
        context = "\n\n".join([doc.page_content for doc in docs])
        
        try:
            client = Groq(api_key=groq_api_key)
        except TypeError:
            import groq
            client = groq.Client(api_key=groq_api_key)
        
        prompt = f"""You are a helpful AI assistant analyzing a document. Answer the question based ONLY on the provided context. Be concise and accurate.

Context from the document:
{context[:3000]}

Question: {query}

Instructions:
- Answer directly and concisely
- Use information ONLY from the context above
- If the answer isn't in the context, say "This information is not found in the document"
- For list questions (like "what are the limitations"), list all items clearly

Answer:"""

        with st.spinner("🚀 Groq AI analyzing... (super fast!)"):
            chat_completion = client.chat.completions.create(
                messages=[{
                    "role": "user",
                    "content": prompt,
                }],
                model="openai/gpt-oss-120b",
                temperature=0.3,
                max_tokens=500,
            )
            
            answer = chat_completion.choices[0].message.content.strip()
            return answer or "This information is not found in the document"

    except Exception as e:
        st.error(f"⚠️ Groq API Error: {str(e)}")
        try:
            docs = vectorstore.similarity_search(query, k=3)
            if docs:
                context = "\n\n".join([doc.page_content[:500] for doc in docs])
                return f"**Relevant Information (Fallback):**\n\n{context}"
        except:
            pass
        return "❌ Unable to find relevant information."

def main():
    st.title("🤖 RAG Q&A App (Powered by Groq)")
    st.markdown("Upload documents or provide links to ask questions!")

    input_type = st.selectbox(
        "Select Input Type",
        ["PDF", "Text", "DOCX", "TXT"]            
    )

    input_data = None

    if input_type == "Text":
        input_data = st.text_area("Enter your text here", height=200)
        
    elif input_type == 'PDF':
        input_data = st.file_uploader("Upload a PDF file", type=["pdf"])
        
    elif input_type == 'TXT':
        input_data = st.file_uploader("Upload a text file", type=['txt'])
        
    elif input_type == 'DOCX':
        input_data = st.file_uploader("Upload a DOCX file", type=['docx', 'doc'])

    if st.button("📚 Process Document", type="primary"):
        try:
            with st.spinner("Processing your document..."):
                vectorstore = process_input(input_type, input_data)
                st.session_state["vectorstore"] = vectorstore
                st.success("Document processed successfully! You can now ask questions.")
        except Exception as e:
            st.error(f"❌ Error: {str(e)}")

    if "vectorstore" in st.session_state:
        st.divider()
        st.subheader("💬 Ask Your Question")
        query = st.text_input("Enter your question here")
        
        if st.button("🔍 Get Answer"):
            if query.strip():
                try:
                    with st.spinner("Finding answer..."):
                        answer = answer_question(st.session_state["vectorstore"], query)
                        st.markdown("### Answer:")
                        st.write(answer)
                except Exception as e:
                    st.error(f"❌ Error: {str(e)}")
            else:
                st.warning("⚠️ Please enter a question")

if __name__ == "__main__":
    main()
